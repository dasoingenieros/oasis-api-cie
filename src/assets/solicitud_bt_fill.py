#!/usr/bin/env python3
"""
solicitud_bt_fill.py — Rellena la Solicitud BT oficial (.docx)

Uso: python3 solicitud_bt_fill.py <template.docx> <output.docx> <output.pdf> <data.json>

JSON structure:
{
  "titular": { "nif":"", "apellido1":"", "apellido2":"", "nombre":"", "email":"",
               "tipoVia":"", "nombreVia":"", "numero":"", "bloque":"", "portal":"",
               "escalera":"", "piso":"", "puerta":"", "localidad":"", "provincia":"",
               "cp":"", "telefono":"", "movil":"" },
  "representante": { ... same fields ... },
  "empresa": { ... same + "categoria":"", "numRegistro":"", "instaladorNombre":"" },
  "proyectista": { ... },
  "directorObra": { ... },
  "emplazamiento": { "tipoVia":"", "nombreVia":"", "numero":"", "cp":"", "localidad":"" },
  "tipoExpediente": { "tipo":"NUEVA|MODIFICACION|AMPLIACION", "numRegistro":"" },
  "tipoInstalacion": { "vivienda":"1", "irve":"1", ... },
  "documentacion": { "mtd":"X", "cie":"X", ... },
  "firma": { "lugar":"MADRID", "dia":"22", "mes":"febrero", "anio":"2026" }
}
"""
import sys
import json
import subprocess
import os
from docx import Document
from docx.shared import Pt


# Default font for filled values — matches the template
FONT_NAME = 'Arial'
FONT_SIZE = Pt(9)


def write_cell(cell, value):
    """Escribe valor preservando formato del párrafo existente, forzando Arial 9pt."""
    if not value:
        return
    text = str(value)
    if cell.paragraphs and cell.paragraphs[0].runs:
        run = cell.paragraphs[0].runs[0]
        run.text = text
        # Ensure font matches template
        if run.font.name is None:
            run.font.name = FONT_NAME
        if run.font.size is None:
            run.font.size = FONT_SIZE
    else:
        # No existing run — create one with proper formatting
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p.text = ''  # Clear any direct text
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE


def set_after_label(row, label_substring, value):
    """Busca celda con label y escribe value en la siguiente celda vacía."""
    if not value:
        return
    cells = row.cells
    for i, cell in enumerate(cells):
        if label_substring.lower() in cell.text.strip().lower():
            for j in range(i + 1, len(cells)):
                if cells[j].text.strip() == '' and cells[j]._tc != cells[i]._tc:
                    if j > 0 and cells[j]._tc == cells[j-1]._tc:
                        continue
                    write_cell(cells[j], value)
                    return
            break


def fill_persona_table(table, data, has_empresa_fields=False):
    """Rellena tabla de datos personales (T0-T4)."""
    rows = table.rows

    # R0: NIF, Apellido1, Apellido2
    write_cell(rows[0].cells[1], data.get('nif', ''))
    set_after_label(rows[0], 'Primer Apellido', data.get('apellido1', ''))
    set_after_label(rows[0], 'Segundo Apellido', data.get('apellido2', ''))

    # R1: Nombre, Email
    set_after_label(rows[1], 'Nombre', data.get('nombre', ''))
    set_after_label(rows[1], 'Correo', data.get('email', ''))

    offset = 0
    if has_empresa_fields:
        set_after_label(rows[2], 'Categoría', data.get('categoria', ''))
        set_after_label(rows[2], 'Registro', data.get('numRegistro', ''))
        set_after_label(rows[3], 'instalador', data.get('instaladorNombre', ''))
        offset = 2

    # Dirección
    r = 2 + offset
    set_after_label(rows[r], 'Tipo de vía', data.get('tipoVia', ''))
    set_after_label(rows[r], 'Nombre vía', data.get('nombreVia', ''))
    set_after_label(rows[r], 'Nº', data.get('numero', ''))

    # Detalle
    r = 3 + offset
    set_after_label(rows[r], 'Bloque', data.get('bloque', ''))
    set_after_label(rows[r], 'Portal', data.get('portal', ''))
    set_after_label(rows[r], 'Escalera', data.get('escalera', ''))
    set_after_label(rows[r], 'Piso', data.get('piso', ''))
    set_after_label(rows[r], 'Puerta', data.get('puerta', ''))
    set_after_label(rows[r], 'Localidad', data.get('localidad', ''))

    # Provincia/CP/Tels
    r = 4 + offset
    set_after_label(rows[r], 'Provincia', data.get('provincia', ''))
    set_after_label(rows[r], 'CP', data.get('cp', ''))
    set_after_label(rows[r], 'Fijo', data.get('telefono', ''))
    set_after_label(rows[r], 'Móvil', data.get('movil', ''))


def fill_emplazamiento(table, data):
    """Tabla 5: solo dirección + CP + localidad."""
    rows = table.rows
    set_after_label(rows[0], 'Tipo de vía', data.get('tipoVia', ''))
    set_after_label(rows[0], 'Nombre vía', data.get('nombreVia', ''))
    set_after_label(rows[0], 'Nº', data.get('numero', ''))
    set_after_label(rows[1], 'CP', data.get('cp', ''))
    set_after_label(rows[1], 'Localidad', data.get('localidad', ''))


def fill_tipo_expediente(table, data):
    """Tabla 6: marcar tipo expediente reemplazando checkbox con X."""
    tipo = data.get('tipo', 'NUEVA').upper()
    cell0 = table.rows[0].cells[0]

    # Map tipo to paragraph index in cell
    tipo_to_paragraph = {
        'NUEVA': 0,
        'MODIFICACION': 1,
        'AMPLIACION': 2,
        'MULTIINSTALACION': 3,
    }
    target_p_idx = tipo_to_paragraph.get(tipo, 0)

    for i, p in enumerate(cell0.paragraphs):
        if i == target_p_idx:
            # Find the empty runs (checkbox characters) and replace with X
            for run in p.runs:
                if run.text.strip() == '' and run.font.size is None:
                    run.text = 'X'
                    run.font.name = FONT_NAME
                    run.font.size = Pt(10)
                    break
            break

    # Fill nº registro existente if provided
    num_reg = data.get('numRegistro', '')
    if num_reg:
        cells = table.rows[0].cells
        seen = set()
        for c in cells:
            tc_id = id(c._tc)
            if tc_id not in seen:
                seen.add(tc_id)
                if c != cell0 and c.text.strip() == '':
                    write_cell(c, num_reg)
                    break



def fill_tipo_instalacion(table, data):
    """Tabla 7: marcar tipo con número."""
    tipo_map = {
        'alumbradoExterior': 1, 'autoconsumo': 2, 'cercasElectricas': 3,
        'caldeo': 4, 'rotulosLuminosos': 5, 'garajes': 6, 'generacion': 7,
        'industrial': 8, 'enlaceComunes': 9, 'irve': 10,
        'riesgoIncendio': 11, 'localEspecial': 12, 'localMojado': 13,
        'localOficina': 14, 'lpcEspectaculos': 15, 'lpcReunion': 16,
        'lpcOtros': 17, 'elevacion': 18, 'piscinas': 19,
        'tensionesEspeciales': 20, 'quirofanos': 21, 'temporal': 22,
        'vivienda': 23, 'otras': 24,
    }
    for key, row_idx in tipo_map.items():
        val = data.get(key, '')
        if val:
            write_cell(table.rows[row_idx].cells[1], str(val))
    total = data.get('totalMulti', '')
    if total:
        write_cell(table.rows[25].cells[2], str(total))


def fill_documentacion(table, data):
    """Tabla 8: marcar con X."""
    doc_map = {
        'tasaDG': 1, 'tasaEICI': 2, 'mtd': 3, 'cie': 4,
        'dossierUsuario': 5, 'proyecto': 6, 'direccionObra': 7,
        'variaciones': 8, 'inspeccionPeriodica': 9, 'contratoMantenimiento': 10,
        'acreditacionEmpresa': 11, 'autoconsumoDoc': 12, 'irveDoc': 13,
        'declaracionInspeccion': 14, 'otros': 15,
    }
    for key, row_idx in doc_map.items():
        val = data.get(key, '')
        if val:
            write_cell(table.rows[row_idx].cells[1], str(val))


def fill_firma(doc, data):
    """Rellena párrafo firma."""
    lugar = data.get('lugar', '')
    dia = data.get('dia', '')
    mes = data.get('mes', '')
    anio = data.get('anio', '')
    firma_text = f"En {lugar}, a {dia} de {mes} de {anio}"
    for p in doc.paragraphs:
        if 'En …' in p.text or 'En ……' in p.text or 'En \u2026' in p.text:
            # Preserve formatting from existing run
            if p.runs:
                for run in p.runs:
                    run.text = ''
                p.runs[0].text = firma_text
            else:
                p.text = firma_text
            break


def convert_to_pdf(docx_path, pdf_path):
    """Convierte .docx a .pdf con LibreOffice."""
    out_dir = os.path.dirname(pdf_path) or '.'
    subprocess.run([
        'soffice', '--headless', '--convert-to', 'pdf',
        '--outdir', out_dir, docx_path
    ], check=True, timeout=30, capture_output=True)
    generated = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + '.pdf')
    if generated != pdf_path and os.path.exists(generated):
        os.rename(generated, pdf_path)


def main():
    if len(sys.argv) < 5:
        print("Usage: solicitud_bt_fill.py <template.docx> <output.docx> <output.pdf> <data.json>")
        sys.exit(1)

    template_path, output_docx, output_pdf, data_json = sys.argv[1:5]

    with open(data_json, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document(template_path)
    tables = doc.tables

    if data.get('titular'):
        fill_persona_table(tables[0], data['titular'])
    if data.get('representante'):
        fill_persona_table(tables[1], data['representante'])
    if data.get('empresa'):
        fill_persona_table(tables[2], data['empresa'], has_empresa_fields=True)
    if data.get('proyectista'):
        fill_persona_table(tables[3], data['proyectista'])
    if data.get('directorObra'):
        fill_persona_table(tables[4], data['directorObra'])
    if data.get('emplazamiento'):
        fill_emplazamiento(tables[5], data['emplazamiento'])
    if data.get('tipoExpediente'):
        fill_tipo_expediente(tables[6], data['tipoExpediente'])
    if data.get('tipoInstalacion'):
        fill_tipo_instalacion(tables[7], data['tipoInstalacion'])
    if data.get('documentacion'):
        fill_documentacion(tables[8], data['documentacion'])
    if data.get('firma'):
        fill_firma(doc, data['firma'])

    doc.save(output_docx)
    print(f"DOCX saved: {output_docx}")

    convert_to_pdf(output_docx, output_pdf)
    print(f"PDF saved: {output_pdf}")


if __name__ == '__main__':
    main()
